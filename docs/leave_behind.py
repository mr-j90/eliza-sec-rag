"""Build the client leave-behind as a .docx.

    uv run --with python-docx python docs/leave_behind.py

Generated rather than hand-written, for the same reason `docs/PROMPT_TEMPLATE.md` is: the
numbers in it come from the system, and a document maintained by hand drifts from the thing it
describes.

Audience is the **client, after the call** — not the presenter. `docs/WALKTHROUGH.md` is the
run-of-show and stays internal. This is business-first, skimmable, and honest about limits,
because a leave-behind that oversells gets read a second time by someone looking for the catch.
"""

from __future__ import annotations

import subprocess
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO = Path(__file__).resolve().parent.parent
LOGO = REPO.parent / "branding" / "689e6c310d474c5ffdda69b2_ChatGPT Image Aug 14, 2025, 04_01_11 PM-456x150.png"
OUT = REPO / "docs" / "eliza-sec-filings-assistant.docx"

AUTHOR = "Jordan Taylor"
ROLE = "Forward Deployed Engineer"

FONT = "Arial"
INK = RGBColor(0x11, 0x11, 0x11)
MUTED = RGBColor(0x66, 0x66, 0x66)
RULE = RGBColor(0xD8, 0xD8, 0xD8)


# --- small helpers ------------------------------------------------------------------------


def style_base(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.22


def para(document, text="", *, size=10.5, bold=False, color=INK, before=0, after=8,
         align=None, italic=False, indent=None):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    if indent is not None:
        p.paragraph_format.left_indent = Inches(indent)
    if text:
        run = p.add_run(text)
        run.font.name = FONT
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = color
    return p


def rich(document, parts, *, size=10.5, before=0, after=8, indent=None):
    """`parts` is a list of (text, bold) so a sentence can emphasise a number."""
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if indent is not None:
        p.paragraph_format.left_indent = Inches(indent)
    for text, bold in parts:
        run = p.add_run(text)
        run.font.name = FONT
        run.font.size = Pt(size)
        run.bold = bold
        run.font.color.rgb = INK
    return p


def heading(document, text, *, level=1):
    sizes = {1: 15, 2: 11.5}
    p = para(
        document, text,
        size=sizes[level], bold=True,
        before=20 if level == 1 else 14,
        after=6 if level == 1 else 4,
    )
    if level == 1:
        _bottom_rule(p)
    return p


def _bottom_rule(paragraph) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "D8D8D8")
    borders.append(bottom)
    pPr.append(borders)


def bullet(document, parts, *, size=10.5):
    p = document.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.28)
    for text, bold in (parts if isinstance(parts, list) else [(parts, False)]):
        run = p.add_run(text)
        run.font.name = FONT
        run.font.size = Pt(size)
        run.bold = bold
        run.font.color.rgb = INK
    return p


def table(document, headers, rows, *, widths=None):
    t = document.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for index, label in enumerate(headers):
        cell = t.rows[0].cells[index]
        cell.text = ""
        run = cell.paragraphs[0].add_run(label)
        run.font.name = FONT
        run.font.size = Pt(9)
        run.bold = True
        run.font.color.rgb = INK
        _shade(cell, "F2F2F2")
    for row in rows:
        cells = t.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = ""
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(str(value))
            run.font.name = FONT
            run.font.size = Pt(9)
            run.font.color.rgb = INK
    if widths:
        for row in t.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    document.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def _shade(cell, hex_colour: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_colour)
    cell._tc.get_or_add_tcPr().append(shd)


def page_number_footer(section) -> None:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("eliza  ·  SEC Filings Research Assistant  ·  ")
    run.font.name = FONT
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)


def git_revision() -> str:
    try:
        out = subprocess.run(
            ["git", "-C", str(REPO), "rev-parse", "--short", "HEAD"],
            capture_output=True, text=True, timeout=5,
        )
        return out.stdout.strip() or "uncommitted"
    except Exception:
        return "uncommitted"


# --- the document -------------------------------------------------------------------------


def build() -> Path:
    document = Document()
    style_base(document)

    section = document.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    page_number_footer(section)

    # ---- cover -------------------------------------------------------------------------
    para(document, after=36)
    if LOGO.is_file():
        logo = document.add_paragraph()
        logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        logo.add_run().add_picture(str(LOGO), width=Inches(1.9))
        logo.paragraph_format.space_after = Pt(40)

    para(document, "SEC Filings Research Assistant", size=26, bold=True, after=4)
    para(
        document,
        "Answering diligence questions from SEC filings, with every claim traceable "
        "to a filing and a period.",
        size=13, color=MUTED, after=44,
    )

    para(document, "Prepared for", size=9, bold=True, color=MUTED, after=2)
    para(document, "Private equity diligence team", size=11, after=16)
    para(document, "Prepared by", size=9, bold=True, color=MUTED, after=2)
    para(document, AUTHOR, size=11, bold=True, after=1)
    para(document, ROLE + ", eliza", size=10, color=MUTED, after=16)
    para(document, "Date", size=9, bold=True, color=MUTED, after=2)
    para(document, date.today().strftime("%-d %B %Y"), size=11, after=30)

    para(
        document,
        f"Corpus: 246 SEC filings (10-K and 10-Q) from 54 US public companies, 2023–2025.  "
        f"Build {git_revision()}.",
        size=8.5, color=MUTED,
    )

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ---- 1. what it is -----------------------------------------------------------------
    heading(document, "What this is")
    para(
        document,
        "A research assistant for SEC filings. You ask a business question in plain "
        "English; it finds the relevant passages across 246 filings and returns a "
        "structured answer in which every claim carries a citation back to a specific "
        "filing, period and section.",
    )
    rich(document, [
        ("It is deliberately a ", False),
        ("research tool, not an oracle", True),
        (". The design goal was not to answer as many questions as possible — it was to "
         "make the system's limits visible, so an analyst always knows what an answer is "
         "standing on. A tool that tells you when it does not know is one you can put in "
         "front of an investment committee.", False),
    ])

    heading(document, "What we demonstrated", level=2)
    table(
        document,
        ["Question you asked", "What it showed"],
        [
            ["Primary risk factors facing Apple, Tesla and JPMorgan, compared",
             "Multi-company retrieval with a guaranteed budget per company, so no company is "
             "crowded out by whichever writes the most vivid prose."],
            ["How NVIDIA's revenue and growth outlook changed over two years",
             "Period-aware retrieval across annual and quarterly filings, with figures that "
             "keep the scale they were reported in."],
            ["Regulatory risks facing major pharmaceutical companies",
             "An industry-level question answered with an explicit statement of how thin the "
             "underlying coverage is."],
            ["A company outside the corpus (Shopify)",
             "A refusal by name, with no findings invented for other companies."],
        ],
        widths=[2.4, 4.1],
    )

    # ---- 2. why trustworthy ------------------------------------------------------------
    heading(document, "Why the answers are trustworthy")
    para(
        document,
        "Four behaviours exist specifically to prevent a confident, well-written, wrong "
        "answer. Each was built because measurement showed the failure was possible.",
        after=10,
    )

    for title, body in [
        ("It refuses questions it cannot support.",
         "Asked about a company with no filings in the corpus, it says so by name and does "
         "not answer for anyone else. Before this rule existed, the same question refused "
         "correctly and then wrote findings for nine other companies."),
        ("Every answer states what it is standing on.",
         "In distinct filings, not passages. For the pharmaceutical question, two companies "
         "have real multi-year coverage and the others have a single filing each — the answer "
         "says so, in the answer, rather than leaving it to be discovered."),
        ("Quarterly risk disclosures are labelled as amendments.",
         "A 10-Q's risk section reports only material changes since the annual report. "
         "Unlabelled, an answer built from one presents an amendment as a complete risk "
         "profile. In the worst case measured, that was a 562-word section standing in for "
         "one ten times its length."),
        ("Citations are verified, not trusted.",
         "Every citation handle is checked against the passages actually retrieved. A handle "
         "that resolves to nothing is flagged on screen rather than quietly removed, because "
         "a citation you cannot check is worse than no citation."),
    ]:
        rich(document, [(title + " ", True), (body, False)], after=7)

    heading(document, "One model call per question", level=2)
    para(
        document,
        "Search, entity resolution, ranking and coverage analysis are all deterministic and "
        "happen before a single language-model call produces the answer. This matters "
        "commercially: cost per question is predictable, and the pipeline can be audited "
        "step by step without re-running the model.",
    )

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ---- 2b. how it works --------------------------------------------------------------
    heading(document, "How a question becomes an answer")
    para(
        document,
        "Two separate processes. The first runs once, over the whole corpus. The second runs "
        "each time a question is asked, and everything in it is deterministic — the language "
        "model is only involved at the final step.",
    )

    heading(document, "Indexing — done once, up front", level=2)
    table(
        document,
        ["Step", "What happens, and why"],
        [
            ["Read and segment",
             "Each filing is split into the sections a reader would recognise — Risk Factors, "
             "Management's Discussion, Financial Statements — and every extracted passage keeps "
             "the company, the period it covers and the section it came from. That metadata is "
             "not bookkeeping: it is how the system later answers “what does Apple say about X” "
             "rather than “something says X”."],
            ["Clean",
             "Tagging residue is removed, paragraph boundaries absent from the source text are "
             "reconstructed, and each table is kept with the caption stating its units and the "
             "row naming its periods. Without this last step a figure can arrive without its "
             "scale."],
            ["Divide into passages",
             "Passages of a few hundred words, cut at natural boundaries rather than at a fixed "
             "length. The average individual risk disclosure in these filings is about the same "
             "size, so a passage tends to correspond to one disclosure rather than to half of "
             "two. The 246 filings produce 30,383 passages."],
            ["Index twice",
             "Every passage is indexed two ways: once for exact wording, once for meaning. Why "
             "both is the next section."],
        ],
        widths=[1.5, 5.0],
    )

    heading(document, "Why two indexes", level=2)
    para(
        document,
        "This is the single most consequential design decision, and it follows from how filings "
        "are written versus how analysts ask.",
        after=6,
    )
    table(
        document,
        ["Search by", "Finds", "Which matters because"],
        [
            ["Exact wording",
             "Named entities, statutes, programmes, tickers — “CHIPS Act”, “Section 232”, "
             "“AAPL”",
             "Filings are dense with precise identifiers, and a near-miss on one of them is a "
             "wrong answer, not a slightly worse one."],
            ["Meaning",
             "Concepts phrased differently from the filing — “supplier concentration” finding "
             "“we depend on a limited number of vendors”",
             "Analysts ask in their own words. A filing almost never uses the phrasing of the "
             "question put to it."],
        ],
        widths=[1.1, 2.6, 2.8],
    )
    rich(document, [
        ("Either alone leaves a visible gap. ", True),
        ("Exact-wording search misses the paraphrase; meaning-based search retrieves passages "
         "that are topically right and about the wrong company. The two result lists are "
         "combined by position rather than by score, so neither has to be calibrated against "
         "the other — a design that stays stable as the corpus grows.", False),
    ], before=4)

    heading(document, "Retrieval — per question", level=2)
    for title, body in [
        ("Understand the question.",
         "Which companies, which periods, annual or quarterly. Rule-based, with no model call, "
         "so the same question always resolves the same way and the result is auditable."),
        ("Search both indexes and combine.",
         "Two searches, one combined ranking."),
        ("Guarantee each company a share.",
         "On a multi-company question, every named company gets its own budget of passages. "
         "Without this, one company's more vivid language crowds the others out — measured "
         "before the rule existed, a three-company question returned fifteen passages for one "
         "company and one for another."),
        ("Remove near-duplicates.",
         "Filings repeat language quarter to quarter. Twenty restatements of one risk is a worse "
         "input than one statement of twenty risks."),
        ("Re-rank the finalists.",
         "A second, more careful model reads the question and each candidate passage together "
         "and re-orders them. It runs locally, adds no per-question cost, and is the largest "
         "single quality gain in the pipeline."),
        ("Assemble and answer.",
         "The surviving passages, each labelled with its source, plus a computed statement of "
         "what the evidence base actually is — then one model call produces the answer."),
    ]:
        rich(document, [(title + " ", True), (body, False)], after=5, indent=0.0)

    para(
        document,
        "The code for all of the above is included in the repository, organised so that each "
        "step is a separate, testable module rather than one pipeline function.",
        before=8,
    )

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ---- 3. value ----------------------------------------------------------------------
    heading(document, "The value: three wrong answers that no longer happen")
    para(
        document,
        "Each of these is an answer an analyst would have acted on, and each is prevented by "
        "a specific mechanism rather than by the model being careful.",
        after=10,
    )
    table(
        document,
        ["The wrong answer", "Why it happened", "What prevents it"],
        [
            ["A figure quoted at the wrong order of magnitude",
             "In these filings the “in millions” caption sits outside the table, on the line "
             "above it. Split the table anywhere below that and the numbers lose their scale.",
             "Captions and period headers are bound to every table passage. Affected passages "
             "fell from 28% to 4%."],
            ["A risk profile that is really one quarter's amendment",
             "A quarterly filing reports only what changed. Read alone, it looks like a "
             "complete picture.",
             "The annual baseline is retrieved alongside it, and quarterly passages are "
             "labelled as amendments."],
            ["An industry conclusion resting on two companies",
             "Four of six pharmaceutical companies in this corpus have a single filing each.",
             "Every answer states its evidence base in distinct filings, per company."],
        ],
        widths=[1.9, 2.4, 2.2],
    )
    rich(document, [
        ("What this is worth. ", True),
        ("None of this makes an analyst faster at reading one filing. It makes 246 filings "
         "searchable with an audit trail — and it makes the system's limits visible, which is "
         "what determines whether the output can be relied on in a deal context.", False),
    ], before=6)

    # ---- 4. measured -------------------------------------------------------------------
    heading(document, "Built on measurement, not assumption")
    para(
        document,
        "Every significant design decision here was driven by measuring this corpus first. "
        "A few of the findings that changed the build:",
        after=10,
    )
    table(
        document,
        ["What we found", "What it changed"],
        [
            ["17.7% of the corpus text was machine-readable tagging residue, not disclosure",
             "Removed before indexing. Left in, it would have polluted search results with "
             "thousands of near-identical fragments."],
            ["40.5% of passages never name their own company",
             "Company, period and section are attached to every passage, because similarity "
             "search alone cannot attribute them."],
            ["The source text has no paragraph structure at all — one filing's risk section "
             "arrives as a single 90,000-character line",
             "Block boundaries are reconstructed. Passages spanning two different report "
             "sections went to zero."],
            ["37 of 246 filings were dated to the wrong year",
             "One corrected derivation. It had also been skewing every “last two years” style "
             "question."],
            ["18 of 54 companies do not end their financial year in December",
             "Citations show the period a filing covers, rather than a fiscal-year label that "
             "can contradict the passage beneath it."],
        ],
        widths=[3.1, 3.4],
    )
    rich(document, [
        ("How quality is measured. ", True),
        ("The system ships with an automated test suite of 266 checks and a retrieval "
         "evaluation harness, viewable in the application itself. We also document why three "
         "commonly-quoted retrieval metrics are ", False),
        ("not", True),
        (" reliable on a corpus like this — knowing why a number misleads is more useful than "
         "the number.", False),
    ], before=6)

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ---- 5. future state ---------------------------------------------------------------
    heading(document, "Where this goes next")
    heading(document, "Nearest-term, highest value", level=2)
    for text in [
        ("Citations that link straight into the filing. Each passage already carries its "
         "source document and section, and each filing carries its public SEC URL — so a "
         "citation becomes a link an analyst can click through to the original text."),
        ("A screen to manage the corpus. Add companies and date ranges and have those filings "
         "pulled in, rather than working from a fixed set."),
        ("Evaluation driven by real questions. Logging the questions analysts actually ask "
         "produces far better test data than anything written in advance."),
        ("Streamed answers, for a faster-feeling interface. The current single response is "
         "deliberate — it lets citations arrive with the text — so this needs designing "
         "alongside, not bolting on."),
    ]:
        bullet(document, text)

    heading(document, "Further out", level=2)
    for text in [
        ("Route numeric questions to the SEC's own structured financial data, so figures come "
         "from a filed data point rather than from reading a table."),
        ("A larger evaluation set with section-level judgements, which is what makes "
         "before-and-after comparisons statistically meaningful rather than directional."),
        ("Self-hosted search models. The filings are public, but the questions asked of them "
         "encode a firm's areas of interest — that is the part worth keeping in-house."),
        ("Multi-tenancy, so several teams can work against the same corpus with separate "
         "entitlements. This is cheapest to design in early."),
    ]:
        bullet(document, text)

    heading(document, "On rebuilding the index", level=2)
    para(
        document,
        "A question worth answering plainly, because the honest answer is more useful than a "
        "reassuring one. Adding new filters, display fields or permissions needs no rebuild. "
        "Changing how documents are divided or which search model is used does require one — "
        "and on a corpus this size that is a background job measured in minutes and cents, "
        "not a project. The architecture is built to make rebuilding cheap rather than to "
        "avoid it.",
    )

    # ---- 6. limits ---------------------------------------------------------------------
    heading(document, "What it does not do today")
    para(
        document,
        "Stated here so none of it is a surprise later.",
        after=8,
    )
    for text in [
        ("15 of 246 filings do not label their sections in a machine-readable way. Their text "
         "is still fully searchable; only the section label is missing."),
        ("Numeric answers are read from tables in the filings rather than from structured "
         "financial data. Figures keep their scale and period, but the structured route would "
         "be more precise."),
        ("The evaluation set is 22 scored questions — enough to catch regressions, not enough "
         "to prove small improvements. Results are reported as directional."),
        ("For an industry-level question, the system reports the companies it used but cannot "
         "tell you which companies it should also have consulted. That needs sector "
         "classification, which this corpus does not carry."),
        ("The corpus is a fixed snapshot. Relative dates are anchored to the newest filing in "
         "it, so “the last two years” stays meaningful as the snapshot ages."),
    ]:
        bullet(document, text)

    # ---- 7. appendix -------------------------------------------------------------------
    heading(document, "Running it yourself")
    para(
        document,
        "The repository is self-contained. Start the search service, build the index once, "
        "then run the application:",
        after=6,
    )
    for line in ["make up", "make index", "make answers", "cd frontend && bun run dev"]:
        mono = para(document, line, size=9.5, after=2, indent=0.28)
        mono.runs[0].font.name = "Consolas"
    para(
        document,
        "A single ready-to-run example request is included, along with the full design record: "
        "the evaluation notes, the prompt-iteration history, and a decision-by-decision log "
        "with the measurement behind each choice.",
        before=8,
    )

    para(document, after=18)
    para(document, f"{AUTHOR} · {ROLE} · eliza", size=9, bold=True, color=MUTED, after=1)
    para(
        document,
        "Prepared as a leave-behind following the demonstration. Figures in this document were "
        "generated from the running system.",
        size=8.5, color=MUTED, italic=True,
    )

    document.core_properties.title = "SEC Filings Research Assistant"
    document.core_properties.author = f"{AUTHOR}, {ROLE}"
    # NB: python-docx exposes no `company` core property — setting one silently creates an
    # unused attribute rather than document metadata. Kept in `category` instead.
    document.core_properties.category = "eliza"
    document.core_properties.comments = "Client leave-behind. Generated by docs/leave_behind.py"

    OUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"wrote {path} ({path.stat().st_size:,} bytes)")
